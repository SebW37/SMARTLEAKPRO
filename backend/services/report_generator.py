"""
Service de génération de rapports
"""

import os
import uuid
from datetime import datetime
from typing import Dict, Any, Optional
from pathlib import Path
import json

from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import pandas as pd

from ..models import Intervention, Rapport, FichierRapport
from ..schemas import GenerationRapport


class ReportGenerator:
    """Générateur de rapports"""
    
    def __init__(self):
        self.templates_dir = Path("backend/templates")
        self.output_dir = Path("backend/static/reports")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Configuration Jinja2
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
    
    async def generate_report(
        self, 
        intervention: Intervention, 
        generation_data: GenerationRapport,
        auteur: str
    ) -> Dict[str, Any]:
        """Générer un rapport complet"""
        
        # Préparer les données du rapport
        report_data = await self._prepare_report_data(intervention, generation_data)
        
        # Générer le contenu selon le format
        if generation_data.format_export.lower() == "pdf":
            file_path = await self._generate_pdf(report_data, generation_data)
        elif generation_data.format_export.lower() == "docx":
            file_path = await self._generate_docx(report_data, generation_data)
        else:
            raise ValueError(f"Format non supporté: {generation_data.format_export}")
        
        # Créer l'enregistrement en base
        rapport = await self._create_rapport_record(
            intervention, 
            generation_data, 
            file_path, 
            auteur,
            report_data
        )
        
        return {
            "rapport_id": str(rapport.id),
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "file_type": generation_data.format_export
        }
    
    async def _prepare_report_data(
        self, 
        intervention: Intervention, 
        generation_data: GenerationRapport
    ) -> Dict[str, Any]:
        """Préparer les données pour le rapport"""
        
        # Charger les inspections liées
        inspections = intervention.inspections if hasattr(intervention, 'inspections') else []
        
        # Charger les fichiers existants
        fichiers = []
        if hasattr(intervention, 'rapports'):
            for rapport in intervention.rapports:
                if hasattr(rapport, 'fichiers'):
                    fichiers.extend(rapport.fichiers)
        
        return {
            "intervention": {
                "id": str(intervention.id),
                "type": intervention.type_intervention.value,
                "statut": intervention.statut.value,
                "date_intervention": intervention.date_intervention,
                "lieu": intervention.lieu,
                "description": intervention.description,
                "technicien": intervention.technicien_assigné,
                "priorite": intervention.priorite,
                "client": {
                    "nom": intervention.client.nom if intervention.client else "N/A",
                    "adresse": intervention.client.adresse if intervention.client else None,
                    "telephone": intervention.client.telephone if intervention.client else None,
                    "email": intervention.client.email if intervention.client else None
                }
            },
            "inspections": [
                {
                    "type": insp.type_inspection,
                    "statut": insp.statut.value,
                    "resultat": insp.resultat,
                    "observations": insp.observations,
                    "date": insp.date_inspection,
                    "coordonnees": insp.coordonnees_gps
                }
                for insp in inspections
            ],
            "fichiers": [
                {
                    "nom": fichier.nom_fichier,
                    "type": fichier.type_fichier,
                    "description": fichier.description,
                    "date_prise": fichier.date_prise,
                    "coordonnees": {
                        "lat": fichier.latitude,
                        "lng": fichier.longitude
                    }
                }
                for fichier in fichiers
            ],
            "metadata": {
                "date_generation": datetime.utcnow(),
                "type_rapport": generation_data.type_rapport.value,
                "template": generation_data.template,
                "inclure_medias": generation_data.inclure_medias,
                "inclure_gps": generation_data.inclure_gps
            }
        }
    
    async def _generate_pdf(self, data: Dict[str, Any], generation_data: GenerationRapport) -> str:
        """Générer un rapport PDF"""
        
        # Charger le template HTML
        template = self.jinja_env.get_template(f"rapport_{generation_data.type_rapport.value}.html")
        html_content = template.render(data=data)
        
        # Générer le PDF
        filename = f"rapport_{uuid.uuid4().hex[:8]}.pdf"
        file_path = self.output_dir / filename
        
        HTML(string=html_content).write_pdf(str(file_path))
        
        return str(file_path)
    
    async def _generate_docx(self, data: Dict[str, Any], generation_data: GenerationRapport) -> str:
        """Générer un rapport DOCX"""
        
        doc = Document()
        
        # Titre
        title = doc.add_heading(f"Rapport d'{generation_data.type_rapport.value.title()}", 0)
        
        # Informations générales
        doc.add_heading("Informations Générales", level=1)
        
        # Tableau des informations
        table_data = [
            ["Client", data["intervention"]["client"]["nom"]],
            ["Type d'intervention", data["intervention"]["type"]],
            ["Date", data["intervention"]["date_intervention"].strftime("%d/%m/%Y %H:%M")],
            ["Lieu", data["intervention"]["lieu"] or "Non spécifié"],
            ["Technicien", data["intervention"]["technicien"] or "Non assigné"],
            ["Priorité", data["intervention"]["priorite"]]
        ]
        
        table = doc.add_table(rows=len(table_data), cols=2)
        table.style = 'Table Grid'
        
        for i, (key, value) in enumerate(table_data):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = str(value)
        
        # Description
        if data["intervention"]["description"]:
            doc.add_heading("Description", level=1)
            doc.add_paragraph(data["intervention"]["description"])
        
        # Inspections
        if data["inspections"]:
            doc.add_heading("Inspections", level=1)
            for i, inspection in enumerate(data["inspections"], 1):
                doc.add_heading(f"Inspection {i}: {inspection['type']}", level=2)
                doc.add_paragraph(f"Statut: {inspection['statut']}")
                if inspection['resultat']:
                    doc.add_paragraph(f"Résultat: {inspection['resultat']}")
                if inspection['observations']:
                    doc.add_paragraph(f"Observations: {inspection['observations']}")
        
        # Fichiers/Médias
        if data["fichiers"] and generation_data.inclure_medias:
            doc.add_heading("Médias Associés", level=1)
            for fichier in data["fichiers"]:
                doc.add_paragraph(f"• {fichier['nom']} ({fichier['type']})")
                if fichier['description']:
                    doc.add_paragraph(f"  Description: {fichier['description']}")
        
        # Coordonnées GPS
        if generation_data.inclure_gps and data["fichiers"]:
            gps_files = [f for f in data["fichiers"] if f["coordonnees"]["lat"]]
            if gps_files:
                doc.add_heading("Coordonnées GPS", level=1)
                for fichier in gps_files:
                    doc.add_paragraph(f"{fichier['nom']}: {fichier['coordonnees']['lat']}, {fichier['coordonnees']['lng']}")
        
        # Pied de page
        doc.add_paragraph(f"Rapport généré le {data['metadata']['date_generation'].strftime('%d/%m/%Y à %H:%M')}")
        
        # Sauvegarder
        filename = f"rapport_{uuid.uuid4().hex[:8]}.docx"
        file_path = self.output_dir / filename
        doc.save(str(file_path))
        
        return str(file_path)
    
    async def _create_rapport_record(
        self,
        intervention: Intervention,
        generation_data: GenerationRapport,
        file_path: str,
        auteur: str,
        report_data: Dict[str, Any]
    ) -> Rapport:
        """Créer l'enregistrement du rapport en base"""
        
        # Cette méthode sera implémentée dans le router
        # pour avoir accès à la session de base de données
        pass


class ReportExporter:
    """Exportateur de rapports en différents formats"""
    
    @staticmethod
    async def export_to_csv(rapports: list, file_path: str):
        """Exporter les rapports en CSV"""
        df = pd.DataFrame(rapports)
        df.to_csv(file_path, index=False, encoding='utf-8')
    
    @staticmethod
    async def export_to_excel(rapports: list, file_path: str):
        """Exporter les rapports en Excel"""
        df = pd.DataFrame(rapports)
        df.to_excel(file_path, index=False, engine='openpyxl')
    
    @staticmethod
    async def export_to_pdf_summary(rapports: list, file_path: str):
        """Exporter un résumé en PDF"""
        doc = SimpleDocTemplate(file_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Titre
        title = Paragraph("Résumé des Rapports", styles['Title'])
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Tableau des rapports
        table_data = [["ID", "Titre", "Type", "Statut", "Date"]]
        for rapport in rapports:
            table_data.append([
                str(rapport.get('id', '')),
                rapport.get('titre', ''),
                rapport.get('type_rapport', ''),
                rapport.get('statut', ''),
                rapport.get('date_creation', '')
            ])
        
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), '#4472C4'),
            ('TEXTCOLOR', (0, 0), (-1, 0), '#FFFFFF'),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), '#F2F2F2'),
            ('GRID', (0, 0), (-1, -1), 1, '#000000')
        ]))
        
        story.append(table)
        doc.build(story)
