"""
Utility functions for report generation.
"""
import os
import tempfile
from django.conf import settings
from django.template.loader import render_to_string
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
import json


def generate_pdf_report(report):
    """Generate PDF report."""
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    temp_file.close()
    
    # Create PDF document
    doc = SimpleDocTemplate(temp_file.name, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph(report.title, title_style))
    story.append(Spacer(1, 12))
    
    # Description
    if report.description:
        story.append(Paragraph(report.description, styles['Normal']))
        story.append(Spacer(1, 12))
    
    # Report content
    content = report.content
    if content:
        # Add content sections
        for section, data in content.items():
            if isinstance(data, dict):
                # Create table for structured data
                table_data = [['Field', 'Value']]
                for key, value in data.items():
                    table_data.append([str(key), str(value)])
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(Paragraph(section, styles['Heading2']))
                story.append(table)
                story.append(Spacer(1, 12))
            else:
                story.append(Paragraph(f"{section}: {data}", styles['Normal']))
                story.append(Spacer(1, 6))
    
    # Build PDF
    doc.build(story)
    
    return temp_file.name


def generate_docx_report(report):
    """Generate Word document report."""
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    temp_file.close()
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading(report.title, 0)
    title.alignment = 1  # Center alignment
    
    # Description
    if report.description:
        doc.add_paragraph(report.description)
        doc.add_paragraph()  # Empty line
    
    # Report content
    content = report.content
    if content:
        for section, data in content.items():
            if isinstance(data, dict):
                # Add section heading
                doc.add_heading(section, level=2)
                
                # Create table for structured data
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Field'
                hdr_cells[1].text = 'Value'
                
                # Data rows
                for key, value in data.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)
                
                doc.add_paragraph()  # Empty line
            else:
                doc.add_heading(section, level=2)
                doc.add_paragraph(str(data))
                doc.add_paragraph()  # Empty line
    
    # Save document
    doc.save(temp_file.name)
    
    return temp_file.name


def generate_html_report(report):
    """Generate HTML report."""
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.html', encoding='utf-8')
    
    # Generate HTML content
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{report.title}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #333; text-align: center; }}
            h2 {{ color: #666; border-bottom: 2px solid #eee; }}
            table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
            .section {{ margin: 20px 0; }}
        </style>
    </head>
    <body>
        <h1>{report.title}</h1>
        {f'<p>{report.description}</p>' if report.description else ''}
    """
    
    # Add content
    content = report.content
    if content:
        for section, data in content.items():
            html_content += f'<div class="section"><h2>{section}</h2>'
            
            if isinstance(data, dict):
                html_content += '<table><tr><th>Field</th><th>Value</th></tr>'
                for key, value in data.items():
                    html_content += f'<tr><td>{key}</td><td>{value}</td></tr>'
                html_content += '</table>'
            else:
                html_content += f'<p>{data}</p>'
            
            html_content += '</div>'
    
    html_content += """
    </body>
    </html>
    """
    
    temp_file.write(html_content)
    temp_file.close()
    
    return temp_file.name


def generate_excel_report(report):
    """Generate Excel report."""
    # Create temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx')
    temp_file.close()
    
    # Create workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Report"
    
    # Title
    ws['A1'] = report.title
    ws['A1'].font = Font(size=16, bold=True)
    ws.merge_cells('A1:D1')
    
    # Description
    if report.description:
        ws['A3'] = report.description
        ws.merge_cells('A3:D3')
        current_row = 4
    else:
        current_row = 3
    
    # Content
    content = report.content
    if content:
        for section, data in content.items():
            # Section header
            ws[f'A{current_row}'] = section
            ws[f'A{current_row}'].font = Font(size=14, bold=True)
            current_row += 1
            
            if isinstance(data, dict):
                # Headers
                ws[f'A{current_row}'] = 'Field'
                ws[f'B{current_row}'] = 'Value'
                ws[f'A{current_row}'].font = Font(bold=True)
                ws[f'B{current_row}'].font = Font(bold=True)
                current_row += 1
                
                # Data
                for key, value in data.items():
                    ws[f'A{current_row}'] = str(key)
                    ws[f'B{current_row}'] = str(value)
                    current_row += 1
            else:
                ws[f'A{current_row}'] = str(data)
                current_row += 1
            
            current_row += 1  # Empty row
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save workbook
    wb.save(temp_file.name)
    
    return temp_file.name
